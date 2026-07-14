import os
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from contextlib import redirect_stdout
from io import StringIO
from pathlib import Path

from PySide6.QtCore import Qt
from PySide6.QtGui import QFont, QFontDatabase, QImage, QPainter, QPageLayout
from PySide6.QtPrintSupport import QPrinter, QPrintDialog
from PySide6.QtSvg import QSvgRenderer
from PySide6.QtWidgets import (
    QListView,
    QApplication,
    QComboBox,
    QFileDialog,
    QFormLayout,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QPlainTextEdit,
    QProgressDialog,
    QSpinBox,
    QTabWidget,
    QVBoxLayout,
    QWidget,
)

import db_report
import import_quiz_results


def detect_office_converter() -> str | None:
    """Return absolute path to soffice/libreoffice if available."""
    for name in ("soffice", "libreoffice"):
        resolved = shutil.which(name)
        if resolved:
            return resolved

    for path in (
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/snap/bin/soffice",
        "/opt/libreoffice/program/soffice",
    ):
        if Path(path).exists():
            return path

    return None


class ImportTab(QWidget):
    def __init__(self):
        super().__init__()
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout(self)

        self.db_path = QLineEdit(os.path.join("data", "quiz_results.db"))
        db_btn = QPushButton("Browse...")
        db_btn.clicked.connect(self._browse_db)

        file_box = QGroupBox("Import single file")
        file_layout = QHBoxLayout(file_box)
        self.file_path = QLineEdit()
        file_browse = QPushButton("Browse...")
        file_browse.clicked.connect(self._browse_file)
        file_import = QPushButton("Import File")
        file_import.clicked.connect(self._import_file)
        file_layout.addWidget(self.file_path)
        file_layout.addWidget(file_browse)
        file_layout.addWidget(file_import)

        folder_box = QGroupBox("Import folder")
        folder_layout = QHBoxLayout(folder_box)
        self.folder_path = QLineEdit()
        folder_browse = QPushButton("Browse...")
        folder_browse.clicked.connect(self._browse_folder)
        folder_import = QPushButton("Import Folder")
        folder_import.clicked.connect(self._import_folder)
        folder_layout.addWidget(self.folder_path)
        folder_layout.addWidget(folder_browse)
        folder_layout.addWidget(folder_import)

        theme_box = QGroupBox("Import event themes")
        theme_layout = QHBoxLayout(theme_box)
        self.theme_file_path = QLineEdit(os.path.join("data", "themes", "PQ_Themen.xlsx"))
        theme_browse = QPushButton("Browse...")
        theme_browse.clicked.connect(self._browse_theme_file)
        theme_import = QPushButton("Import Themes")
        theme_import.clicked.connect(self._import_themes)
        theme_layout.addWidget(self.theme_file_path)
        theme_layout.addWidget(theme_browse)
        theme_layout.addWidget(theme_import)

        db_row = QHBoxLayout()
        db_row.addWidget(QLabel("Database:"))
        db_row.addWidget(self.db_path)
        db_row.addWidget(db_btn)

        self.log = QPlainTextEdit()
        self.log.setReadOnly(True)

        layout.addLayout(db_row)
        layout.addWidget(file_box)
        layout.addWidget(folder_box)
        layout.addWidget(theme_box)
        layout.addWidget(QLabel("Import log"))
        layout.addWidget(self.log)

    def _browse_db(self):
        path, _ = QFileDialog.getSaveFileName(self, "Select database", self.db_path.text(), "SQLite DB (*.db)")
        if path:
            self.db_path.setText(path)

    def _browse_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Select quiz file",
            "",
            "Excel Files (*.xlsx *.xlsm *.xltx *.xltm)",
        )
        if path:
            self.file_path.setText(path)

    def _browse_folder(self):
        path = QFileDialog.getExistingDirectory(self, "Select quiz folder")
        if path:
            self.folder_path.setText(path)

    def _browse_theme_file(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Select themes file",
            self.theme_file_path.text().strip() or "",
            "Excel Files (*.xlsx *.xlsm *.xltx *.xltm)",
        )
        if path:
            self.theme_file_path.setText(path)

    def _import_file(self):
        path = self.file_path.text().strip()
        db = self.db_path.text().strip()
        if not path:
            QMessageBox.warning(self, "Missing input", "Please select a file.")
            return

        try:
            result = import_quiz_results.import_quiz_file(path, db)
            self.log.appendPlainText(
                f"Imported {result['teams_imported']} teams from {result['event_date']} @ {result['location']}"
            )
        except Exception as exc:
            self.log.appendPlainText(f"ERROR: {exc}")
            QMessageBox.critical(self, "Import failed", str(exc))

    def _import_folder(self):
        path = self.folder_path.text().strip()
        db = self.db_path.text().strip()
        if not path:
            QMessageBox.warning(self, "Missing input", "Please select a folder.")
            return

        try:
            summary = import_quiz_results.import_quiz_folder(path, db)
            for item in summary["imports"]:
                self.log.appendPlainText(
                    f"Imported {item['teams_imported']} teams from {item['event_date']} @ {item['location']}"
                )

            if summary.get("errors"):
                self.log.appendPlainText("\nErrors during import:")
                for err in summary["errors"]:
                    self.log.appendPlainText(f"[{err['file']}]\n{err['message']}")

            self.log.appendPlainText(
                f"\nImported {summary['files_imported']} files. Failed: {summary.get('files_failed', 0)}"
            )
        except Exception as exc:
            self.log.appendPlainText(f"ERROR: {exc}")
            QMessageBox.critical(self, "Import failed", str(exc))

    def _import_themes(self):
        path = self.theme_file_path.text().strip()
        db = self.db_path.text().strip()
        if not path:
            QMessageBox.warning(self, "Missing input", "Please select a themes file.")
            return

        try:
            summary = import_quiz_results.import_event_themes(path, db)
            self.log.appendPlainText(
                f"Imported themes: {summary['rows_imported']}/{summary['rows_found']} rows"
            )
            if summary.get("errors"):
                self.log.appendPlainText("Theme import issues:")
                for err in summary["errors"]:
                    self.log.appendPlainText(f"Row {err['row']}: {err['message']}")

            if summary.get("rows_failed", 0) > 0:
                QMessageBox.warning(
                    self,
                    "Theme import completed with issues",
                    (
                        f"Imported {summary['rows_imported']} of {summary['rows_found']} rows. "
                        f"Failed: {summary['rows_failed']}"
                    ),
                )
            else:
                QMessageBox.information(
                    self,
                    "Theme import finished",
                    f"Imported {summary['rows_imported']} rows successfully.",
                )
        except Exception as exc:
            self.log.appendPlainText(f"ERROR: {exc}")
            QMessageBox.critical(self, "Theme import failed", str(exc))


class ReportsTab(QWidget):
    def __init__(self):
        super().__init__()
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout(self)

        controls = QFormLayout()

        self.db_path = QLineEdit(os.path.join("data", "quiz_results.db"))

        self.report_type = QComboBox()
        self.report_type.addItems(
            [
                "list",
                "result",
                "Event Result Team",
                "standings",
                "standing-progress",
                "leaders",
                "team",
                "averages",
                "radar",
                "bonus-efficiency",
                "consistency",
                "difficulty",
                "event-difficulty",
                "round-strength",
            ]
        )

        self.year = QSpinBox()
        self.year.setRange(2000, 2100)
        self.year.setValue(2026)

        self.team = QComboBox()
        self.team.setMaxVisibleItems(14)
        team_view = QListView(self.team)
        team_view.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.team.setView(team_view)
        self.team.setEditable(True)
        self.team.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.team.completer().setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self.team.completer().setFilterMode(Qt.MatchFlag.MatchContains)
        self.round_name = QLineEdit()

        self.min_events = QSpinBox()
        self.min_events.setRange(1, 100)
        self.min_events.setValue(2)

        self.top = QSpinBox()
        self.top.setRange(1, 100)
        self.top.setValue(10)

        self.event_id = QSpinBox()
        self.event_id.setRange(0, 999999)
        self.event_id.setValue(0)

        self.source_file = QLineEdit()
        self.event_date = QLineEdit()
        self.location = QLineEdit()

        event_pick_row = QHBoxLayout()
        self.event_picker = QComboBox()
        self.event_picker.addItem("Select event...", None)
        refresh_events_btn = QPushButton("Refresh events")
        refresh_events_btn.clicked.connect(self._refresh_events)
        apply_event_btn = QPushButton("Use selected event")
        apply_event_btn.clicked.connect(self._apply_selected_event)
        event_pick_row.addWidget(self.event_picker)
        event_pick_row.addWidget(refresh_events_btn)
        event_pick_row.addWidget(apply_event_btn)

        controls.addRow("Database", self.db_path)
        controls.addRow("Report", self.report_type)
        controls.addRow("Year", self.year)
        controls.addRow("Team", self.team)
        controls.addRow("Round name", self.round_name)
        controls.addRow("Min events", self.min_events)
        controls.addRow("Top", self.top)
        controls.addRow("Event ID (0=off)", self.event_id)
        controls.addRow("Source file", self.source_file)
        controls.addRow("Event date (YYYY-MM-DD)", self.event_date)
        controls.addRow("Location", self.location)
        controls.addRow("Event helper", event_pick_row)

        btn_row = QHBoxLayout()
        run_btn = QPushButton("Run report")
        run_btn.clicked.connect(self._run_report)
        pdf_btn = QPushButton("Generate Text PDF")
        pdf_btn.clicked.connect(self._generate_pdf)
        team_pdf_btn = QPushButton("Generate Team PDF")
        team_pdf_btn.clicked.connect(self._generate_team_pdf)
        all_team_pdf_btn = QPushButton("Generate All Team Reports")
        all_team_pdf_btn.clicked.connect(self._generate_all_team_reports)
        btn_row.addWidget(run_btn)
        btn_row.addWidget(pdf_btn)
        btn_row.addWidget(team_pdf_btn)
        btn_row.addWidget(all_team_pdf_btn)

        self.soffice_info = QLabel()
        self.soffice_info.setWordWrap(True)

        self.output = QPlainTextEdit()
        self.output.setReadOnly(True)
        mono_font = QFontDatabase.systemFont(QFontDatabase.SystemFont.FixedFont)
        mono_font.setPointSize(10)
        self.output.setFont(mono_font)

        layout.addLayout(controls)
        layout.addLayout(btn_row)
        layout.addWidget(self.soffice_info)
        layout.addWidget(self.output)

        self.db_path.editingFinished.connect(self._refresh_lookup_data)
        self.year.valueChanged.connect(self._refresh_teams)

        self._refresh_converter_info()
        self._refresh_lookup_data()

    def _selected_team_name(self) -> str:
        selected = self.team.currentData()
        if selected is None:
            return self.team.currentText().strip()
        return str(selected).strip()

    def _refresh_lookup_data(self):
        self._refresh_events()
        self._refresh_teams()

    def _refresh_converter_info(self):
        converter = detect_office_converter()
        if converter:
            self.soffice_info.setText(f"DOCX->PDF converter found: {converter}")
            return

        self.soffice_info.setText(
            "DOCX->PDF converter not found. Install LibreOffice and ensure 'soffice' is available in PATH."
        )

    def _refresh_events(self):
        db = self.db_path.text().strip()
        try:
            events = db_report.get_event_list(db)
        except Exception as exc:
            self.output.setPlainText(f"ERROR while loading events:\n{exc}")
            return

        self.event_picker.clear()
        self.event_picker.addItem("Select event...", None)
        for event in events:
            label = f"#{event['id']} | {event['event_date']} | {event['location']} | {event.get('winner') or 'N/A'}"
            self.event_picker.addItem(label, event)

    def _refresh_teams(self):
        db = self.db_path.text().strip()
        selected_team = self._selected_team_name()

        self.team.clear()
        self.team.addItem("Select team...", "")

        try:
            standings = db_report.get_championship_standings(self.year.value(), db).get("standings", [])
        except Exception:
            return

        team_names = []
        for row in standings:
            team_name = str(row.get("team_name", "")).strip()
            if team_name:
                team_names.append(team_name)

        for team_name in sorted(team_names, key=str.casefold):
            self.team.addItem(team_name, team_name)

        if selected_team:
            idx = self.team.findData(selected_team)
            if idx >= 0:
                self.team.setCurrentIndex(idx)

    def _apply_selected_event(self):
        data = self.event_picker.currentData()
        if not data:
            return
        self.event_id.setValue(int(data["id"]))
        self.event_date.setText(str(data["event_date"]))
        self.location.setText(str(data["location"]))

        # Keep source_file in sync when available.
        try:
            event_detail = db_report.get_event_result(
                db_path=self.db_path.text().strip(),
                event_id=int(data["id"]),
            )
            self.source_file.setText(str(event_detail["event"]["source_file"]))
        except Exception:
            # Non-fatal: user can still use id/date/location.
            pass

    def _capture(self, fn, *args, **kwargs):
        buf = StringIO()
        with redirect_stdout(buf):
            fn(*args, **kwargs)
        return buf.getvalue()

    def _generate_pdf(self):
        text = self.output.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "No report", "Run a report first before generating a PDF.")
            return

        path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "", "PDF Files (*.pdf)")
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            path += ".pdf"

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(path)
        printer.setPageOrientation(QPageLayout.Orientation.Portrait)

        self.output.print_(printer)
        QMessageBox.information(self, "PDF saved", f"Report saved to:\n{path}")

    def _team_pdf_default_name(self, team_name: str, year: int) -> str:
        safe_team = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in team_name.strip()).strip("_")
        if not safe_team:
            safe_team = "team"
        return f"{safe_team}_{year}_team_report.pdf"

    def _template_values(self, report: dict) -> dict:
        events = report.get("season_events", [])
        radar_rounds = report.get("radar_rounds", [])
        best_result = report.get("best_result")

        best_pos_event = None
        positioned_events = [event for event in events if event.get("position") is not None]
        if positioned_events:
            best_pos_event = min(positioned_events, key=lambda event: event["position"])

        ranked_rounds = [row for row in radar_rounds if row.get("position") is not None]
        best_category = None
        if ranked_rounds:
            best_category = min(
                ranked_rounds,
                key=lambda row: (
                    row["position"],
                    -(row["avg_points"] if row.get("avg_points") is not None else 0),
                    row["round_name"],
                ),
            )

        year_events_count = report.get("year_events_count")
        if year_events_count is None:
            year_events_count = len(events)

        first_places = sum(1 for event in events if event.get("position") == 1)
        second_places = sum(1 for event in events if event.get("position") == 2)
        third_places = sum(1 for event in events if event.get("position") == 3)

        return {
            "Date": datetime.now().strftime("%d.%m.%Y"),
            "YEAR": str(report["year"]),
            "Team Name": report["team_name"],
            "Teams": str(report.get("teams_total", "-")),
            "ChampionshipPoints": "-" if report["championship_points"] is None else str(report["championship_points"]),
            "ChampionshipPosition": "-" if report["championship_place"] is None else str(report["championship_place"]),
            "BestPos": "-" if best_pos_event is None else str(best_pos_event["position"]),
            "DateBestPos": "-" if best_pos_event is None else best_pos_event["event_date"],
            "LocBestPos": "-" if best_pos_event is None else best_pos_event["location"],
            "BestPts": "-" if best_result is None else str(best_result["total_points"]),
            "DateBestPts": "-" if best_result is None else best_result["event_date"],
            "LocBestPts": "-" if best_result is None else best_result["location"],
            "AvgPoints": "-" if report["average_points"] is None else f"{report['average_points']:.2f}",
            "BonusEff": "-" if report.get("bonus_efficiency_avg") is None else f"{report['bonus_efficiency_avg'] * 100:.1f}%",
            "BestCat": "-" if best_category is None else best_category["round_name"],
            "BestCatPos": "-" if best_category is None else str(best_category["position"]),
            "nTeams": "-" if best_category is None else str(best_category["total_teams"]),
            "Participations": str(report["participation_count"]),
            "nRoundsPlayed": str(year_events_count),
            "1": str(first_places),
            "2": str(second_places),
            "3": str(third_places),
        }

    def _iter_docx_paragraphs(self, document):
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

        # Include placeholders in headers/footers (e.g. footer date tokens).
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

            for paragraph in section.footer.paragraphs:
                yield paragraph
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

    def _replace_placeholder_in_paragraph(self, paragraph, placeholder: str, value: str) -> bool:
        run_texts = [run.text for run in paragraph.runs]
        full_text = "".join(run_texts)
        if placeholder not in full_text:
            return False

        start = full_text.find(placeholder)
        end = start + len(placeholder)

        boundaries = []
        offset = 0
        for idx, text in enumerate(run_texts):
            boundaries.append((idx, offset, offset + len(text)))
            offset += len(text)

        start_idx = None
        end_idx = None
        for idx, b_start, b_end in boundaries:
            if start_idx is None and b_start <= start < b_end:
                start_idx = idx
            if b_start < end <= b_end:
                end_idx = idx
                break

        if start_idx is None or end_idx is None:
            paragraph.text = full_text.replace(placeholder, value)
            return True

        start_run = paragraph.runs[start_idx]
        end_run = paragraph.runs[end_idx]

        start_local = start - boundaries[start_idx][1]
        end_local = end - boundaries[end_idx][1]

        prefix = start_run.text[:start_local]
        suffix = end_run.text[end_local:]
        start_run.text = prefix + value + suffix

        for idx in range(start_idx + 1, end_idx + 1):
            paragraph.runs[idx].text = ""

        return True

    def _render_radar_png(self, svg_path: str, png_path: str, width: int = 1800, height: int = 1200) -> str:
        renderer = QSvgRenderer(svg_path)
        if not renderer.isValid():
            raise RuntimeError("Radar SVG could not be loaded.")

        image = QImage(width, height, QImage.Format.Format_ARGB32)
        image.fill(Qt.GlobalColor.white)

        painter = QPainter(image)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
        painter.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform, True)
        renderer.render(painter)
        painter.end()

        if not image.save(png_path, "PNG"):
            raise RuntimeError("Failed to render radar PNG.")
        return png_path

    def _replace_docx_placeholders(self, document, values: dict) -> None:
        for placeholder, value in values.items():
            token = "{" + placeholder + "}"
            replacement = str(value)
            for paragraph in self._iter_docx_paragraphs(document):
                while self._replace_placeholder_in_paragraph(paragraph, token, replacement):
                    continue

            # Also replace placeholders inside text boxes/shapes where python-docx
            # does not expose paragraphs through document.paragraphs/tables.
            self._replace_placeholder_in_xml_text_nodes(document.part.element, token, replacement)
            for section in document.sections:
                self._replace_placeholder_in_xml_text_nodes(section.header.part.element, token, replacement)
                self._replace_placeholder_in_xml_text_nodes(section.footer.part.element, token, replacement)

    def _replace_placeholder_in_xml_text_nodes(self, root_element, placeholder: str, value: str) -> None:
        # Replace within WordprocessingML paragraphs (handles split runs like "{" + "2}").
        for paragraph in root_element.xpath(".//w:p"):
            text_nodes = paragraph.xpath(".//w:t")
            if not text_nodes:
                continue
            while self._replace_placeholder_across_text_nodes(text_nodes, placeholder, value):
                pass

        # Also handle DrawingML text (e.g. grouped shapes/text boxes using a:t).
        for paragraph in root_element.xpath(".//a:p"):
            text_nodes = paragraph.xpath(".//a:t")
            if not text_nodes:
                continue
            while self._replace_placeholder_across_text_nodes(text_nodes, placeholder, value):
                pass

    def _replace_placeholder_across_text_nodes(self, text_nodes, placeholder: str, value: str) -> bool:
        """Replace one placeholder occurrence across XML text nodes while preserving layout."""
        run_texts = [node.text or "" for node in text_nodes]
        full_text = "".join(run_texts)
        start = full_text.find(placeholder)
        if start < 0:
            return False
        end = start + len(placeholder)

        boundaries = []
        offset = 0
        for idx, text in enumerate(run_texts):
            boundaries.append((idx, offset, offset + len(text)))
            offset += len(text)

        start_idx = None
        end_idx = None
        for idx, b_start, b_end in boundaries:
            if start_idx is None and b_start <= start < b_end:
                start_idx = idx
            if b_start < end <= b_end:
                end_idx = idx
                break

        if start_idx is None or end_idx is None:
            return False

        start_node = text_nodes[start_idx]
        end_node = text_nodes[end_idx]
        start_local = start - boundaries[start_idx][1]
        end_local = end - boundaries[end_idx][1]

        prefix = (start_node.text or "")[:start_local]
        suffix = (end_node.text or "")[end_local:]
        start_node.text = prefix + value + suffix

        for idx in range(start_idx + 1, end_idx + 1):
            text_nodes[idx].text = ""

        return True

    def _replace_image_in_cell(self, cell, image_path: str) -> bool:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Emu, Inches

        preserved_width = None
        preserved_height = None
        preserved_alignment = None

        # Reuse dimensions of the existing image in this cell, if present.
        for paragraph in cell.paragraphs:
            if preserved_alignment is None:
                preserved_alignment = paragraph.alignment

            for run in paragraph.runs:
                extents = run._r.xpath(".//wp:inline/wp:extent")
                if not extents:
                    extents = run._r.xpath(".//wp:anchor/wp:extent")
                if extents:
                    extent = extents[0]
                    cx = extent.get("cx")
                    cy = extent.get("cy")
                    if cx and cy:
                        preserved_width = Emu(int(cx))
                        preserved_height = Emu(int(cy))
                        break
            if preserved_width is not None and preserved_height is not None:
                break

        cell.text = ""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = preserved_alignment or WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        if preserved_width is not None and preserved_height is not None:
            run.add_picture(image_path, width=preserved_width, height=preserved_height)
        elif cell.width:
            run.add_picture(image_path, width=int(cell.width * 0.92))
        else:
            run.add_picture(image_path, width=Inches(5.8))

        return True

    def _insert_charts_into_docx(self, document, avg_radar_png_path: str, pos_radar_png_path: str, event_chart_png_path: str) -> bool:
        if len(document.tables) < 2:
            return False

        table = document.tables[1]
        if len(table.rows) < 1 or len(table.columns) < 2:
            return False

        left_ok = self._replace_image_in_cell(table.cell(0, 0), avg_radar_png_path)
        right_ok = self._replace_image_in_cell(table.cell(0, 1), pos_radar_png_path)
        bottom_ok = False
        if len(table.rows) >= 2:
            bottom_ok = self._replace_image_in_cell(table.cell(1, 0), event_chart_png_path)

        return left_ok and right_ok and bottom_ok

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        converter = detect_office_converter()
        if converter is None:
            raise RuntimeError(
                "No DOCX-to-PDF converter found. Install LibreOffice and ensure 'soffice' is available in PATH."
            )

        out_dir = Path(pdf_path).parent.resolve()
        out_dir.mkdir(parents=True, exist_ok=True)

        command = [
            converter,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(Path(docx_path).resolve()),
        ]
        result = subprocess.run(command, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(f"DOCX conversion failed: {result.stderr.strip() or result.stdout.strip()}")

        converted = out_dir / (Path(docx_path).stem + ".pdf")
        if not converted.exists():
            raise RuntimeError("DOCX conversion did not produce a PDF file.")

        target = Path(pdf_path).resolve()
        if converted.resolve() != target:
            converted.replace(target)

    def _render_team_pdf_from_docx_template(self, report: dict, output_pdf_path: str) -> None:
        try:
            from docx import Document
        except Exception as exc:
            raise RuntimeError("python-docx is required. Install it with: pip install python-docx") from exc

        template_path = Path("data") / "templates" / "TeamReportTemplate.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        values = self._template_values(report)

        with tempfile.TemporaryDirectory(prefix="team-report-") as tmp_dir:
            tmp_docx = Path(tmp_dir) / "team_report_filled.docx"
            tmp_avg_png = Path(tmp_dir) / "radar_plot_avg.png"
            tmp_pos_png = Path(tmp_dir) / "radar_plot_pos.png"
            tmp_event_png = Path(tmp_dir) / "event_chart.png"

            self._render_radar_png(report["radar_svg_path"], str(tmp_avg_png))
            self._render_radar_png(report["radar_position_svg_path"], str(tmp_pos_png))
            self._render_radar_png(report["event_chart_svg_path"], str(tmp_event_png), width=2200, height=1000)

            document = Document(str(template_path))
            self._replace_docx_placeholders(document, values)
            inserted = self._insert_charts_into_docx(document, str(tmp_avg_png), str(tmp_pos_png), str(tmp_event_png))
            if not inserted:
                raise RuntimeError("Could not place charts in second table.")

            document.save(str(tmp_docx))
            self._convert_docx_to_pdf(str(tmp_docx), output_pdf_path)

    def _generate_team_pdf(self):
        team_name = self._selected_team_name()
        if not team_name:
            QMessageBox.warning(self, "Missing team", "Please select a team before generating the PDF.")
            return

        db = self.db_path.text().strip()
        default_name = self._team_pdf_default_name(team_name, self.year.value())
        path, _ = QFileDialog.getSaveFileName(self, "Save team PDF", default_name, "PDF Files (*.pdf)")
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            path += ".pdf"

        radar_path = None
        radar_position_path = None
        event_chart_path = None
        try:
            report = db_report.get_team_profile_report(team_name, self.year.value(), self.min_events.value(), db)
            radar_path = report.get("radar_svg_path")
            radar_position_path = report.get("radar_position_svg_path")
            event_chart_path = report.get("event_chart_svg_path")

            self._render_team_pdf_from_docx_template(report, path)
            QMessageBox.information(self, "PDF saved", f"Team report saved to:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Team PDF failed", str(exc))
            self.output.setPlainText(f"ERROR:\n{exc}")
        finally:
            if radar_path:
                try:
                    Path(radar_path).unlink(missing_ok=True)
                except Exception:
                    pass
            if radar_position_path:
                try:
                    Path(radar_position_path).unlink(missing_ok=True)
                except Exception:
                    pass
            if event_chart_path:
                try:
                    Path(event_chart_path).unlink(missing_ok=True)
                except Exception:
                    pass

    def _generate_all_team_reports(self):
        db = self.db_path.text().strip()
        year = self.year.value()
        folder = QFileDialog.getExistingDirectory(self, "Select output folder")
        if not folder:
            return

        try:
            standings = db_report.get_championship_standings(year, db)["standings"]
        except Exception as exc:
            QMessageBox.critical(self, "Batch export failed", str(exc))
            return

        if not standings:
            QMessageBox.information(self, "No teams", f"No teams found for {year}.")
            return

        self.output.setPlainText(f"Generating team reports for {year}...\n")
        progress = QProgressDialog(f"Generating team reports for {year}...", None, 0, len(standings), self)
        progress.setWindowTitle("Generating team reports")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setCancelButton(None)
        progress.setMinimumDuration(0)
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        successes = 0
        failures = []

        for index, team in enumerate(standings, start=1):
            team_name = team["team_name"]
            progress.setLabelText(f"Generating {team_name} ({index}/{len(standings)})")
            safe_team = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in team_name.strip()).strip("_")
            if not safe_team:
                safe_team = "team"
            output_path = Path(folder) / f"{safe_team}_{year}_team_report.pdf"

            try:
                report = db_report.get_team_profile_report(team_name, year, self.min_events.value(), db)
                self._render_team_pdf_from_docx_template(report, str(output_path))
                successes += 1
                self.output.appendPlainText(f"[{index}/{len(standings)}] OK: {team_name} -> {output_path.name}")
            except Exception as exc:
                failures.append((team_name, str(exc)))
                self.output.appendPlainText(f"[{index}/{len(standings)}] FAIL: {team_name}: {exc}")

            progress.setValue(index)
            QApplication.processEvents()

        progress.setValue(len(standings))
        QApplication.processEvents()
        progress.close()

        summary = f"Finished. Successes: {successes}. Failures: {len(failures)}."
        if failures:
            summary += "\n\nFailures:\n" + "\n".join(f"- {name}: {message}" for name, message in failures)

        QMessageBox.information(self, "Batch export finished", summary)

    def _run_report(self):
        report = self.report_type.currentText()
        db = self.db_path.text().strip()

        kwargs_event = {
            "db_path": db,
            "event_id": self.event_id.value() if self.event_id.value() > 0 else None,
            "source_file": self.source_file.text().strip() or None,
            "event_date": self.event_date.text().strip() or None,
            "location": self.location.text().strip() or None,
        }

        try:
            if report == "list":
                text = self._capture(db_report.print_event_list, db)
            elif report == "result":
                text = self._capture(db_report.print_event_results, **kwargs_event)
            elif report == "Event Result Team":
                text = self._capture(
                    db_report.print_event_team_result,
                    self._selected_team_name(),
                    **kwargs_event,
                )
            elif report == "standings":
                text = self._capture(db_report.print_championship_standings, self.year.value(), db)
            elif report == "standing-progress":
                text = self._capture(
                    db_report.print_standing_progress_report,
                    self.year.value(),
                    self.top.value(),
                    None,
                    db,
                )
            elif report == "leaders":
                text = self._capture(db_report.print_leaders_report, self.year.value(), self.min_events.value(), db)
            elif report == "team":
                text = self._capture(db_report.print_team_season_results, self._selected_team_name(), self.year.value(), db)
            elif report == "averages":
                text = self._capture(db_report.print_team_round_averages, self._selected_team_name(), self.year.value(), db)
            elif report == "radar":
                text = self._capture(
                    db_report.print_team_radar_report,
                    self._selected_team_name(),
                    self.year.value(),
                    self.min_events.value(),
                    None,
                    db,
                )
            elif report == "bonus-efficiency":
                text = self._capture(
                    db_report.print_team_bonus_efficiency_report,
                    self._selected_team_name(),
                    self.year.value(),
                    db,
                )
            elif report == "consistency":
                text = self._capture(db_report.print_consistency_report, self.year.value(), self.min_events.value(), db)
            elif report == "difficulty":
                text = self._capture(db_report.print_round_difficulty_report, self.year.value(), db)
            elif report == "event-difficulty":
                text = self._capture(db_report.print_event_difficulty_report, **kwargs_event)
            elif report == "round-strength":
                round_name = self.round_name.text().strip() or None
                team_name = self._selected_team_name() or None
                text = self._capture(
                    db_report.print_round_strength_ranking,
                    self.year.value(),
                    round_name,
                    self.min_events.value(),
                    self.top.value(),
                    team_name,
                    db,
                )
            else:
                text = "Unknown report"

            self.output.setPlainText(text)
        except Exception as exc:
            self.output.setPlainText(f"ERROR:\n{exc}")
            QMessageBox.critical(self, "Report failed", str(exc))


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PubQuizStats")
        self.resize(1100, 750)

        tabs = QTabWidget()
        tabs.addTab(ImportTab(), "Import")
        tabs.addTab(ReportsTab(), "Reports")

        self.setCentralWidget(tabs)


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    w = MainWindow()
    w.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
